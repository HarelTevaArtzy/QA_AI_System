from __future__ import annotations

import os
import re
import uuid
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


TEST_DB_FILE = Path(__file__).resolve().parent / f"test-{uuid.uuid4().hex}.db"
os.environ["DATABASE_FILE"] = str(TEST_DB_FILE)
os.environ["AGNO_PROVIDER"] = "disabled"
os.environ["DEFAULT_ADMIN_USERNAME"] = "admin"
os.environ["DEFAULT_ADMIN_PASSWORD"] = "AdminPass123!"

from fastapi.testclient import TestClient

from backend.database import engine
from backend.main import app
from backend.services.agno_agents.scenario_agent import parse_scenario_suggestions


def teardown_module() -> None:
    engine.dispose()
    if TEST_DB_FILE.exists():
        TEST_DB_FILE.unlink()


def login(client: TestClient, username: str = "admin", password: str = "AdminPass123!") -> dict[str, str]:
    response = client.post("/auth/login", json={"username": username, "password": password})
    assert response.status_code == 200
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def create_user(client: TestClient, headers: dict[str, str], username: str, role: str) -> dict:
    response = client.post(
        "/users",
        headers=headers,
        json={
            "username": username,
            "password": "Password123!",
            "role": role,
        },
    )
    assert response.status_code == 201
    return response.json()


def test_scenario_crud_and_exports() -> None:
    with TestClient(app) as client:
        headers = login(client)
        requirement = client.post(
            "/requirements",
            headers=headers,
            json={
                "title": "Prevent duplicate checkout charges",
                "description": "The checkout flow must reject duplicate payment submissions.",
                "parent_id": None,
            },
        )
        assert requirement.status_code == 201
        requirement_id = requirement.json()["id"]

        critical_payload = {
            "title": "Checkout blocks duplicate payment",
            "description": "Regression coverage for duplicate payment retries.",
            "steps": "Open checkout\nSubmit payment\nClick submit again",
            "expected_result": "Second submission is blocked\nNo duplicate charge occurs",
            "priority": "critical",
            "requirement_ids": [requirement_id],
        }
        high_payload = {
            **critical_payload,
            "title": "Login redirect regression",
            "priority": "high",
            "requirement_ids": [],
        }
        low_payload = {
            **critical_payload,
            "title": "Minor UI typo",
            "priority": "low",
            "requirement_ids": [],
        }

        created = client.post("/scenarios", headers=headers, json=critical_payload)
        assert created.status_code == 201
        assert created.json()["requirements"][0]["id"] == requirement_id
        scenario_id = created.json()["id"]

        listed = client.get("/scenarios", headers=headers)
        assert listed.status_code == 200
        assert len(listed.json()) == 1

        higher_priority = client.post("/scenarios", headers=headers, json=high_payload)
        assert higher_priority.status_code == 201

        lower_priority = client.post("/scenarios", headers=headers, json=low_payload)
        assert lower_priority.status_code == 201

        ordered = client.get("/scenarios", headers=headers)
        assert ordered.status_code == 200
        ordered_payload = ordered.json()
        assert ordered_payload[0]["priority"] == "critical"
        assert ordered_payload[1]["priority"] == "high"
        assert ordered_payload[-1]["priority"] == "low"

        updated = client.put(
            f"/scenarios/{scenario_id}",
            headers=headers,
            json=critical_payload,
        )
        assert updated.status_code == 200
        assert updated.json()["priority"] == "critical"

        excel = client.get("/export/scenarios/excel", headers=headers)
        assert excel.status_code == 200
        assert excel.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        workbook = load_workbook(BytesIO(excel.content))
        sheet = workbook.active
        excel_priorities = [sheet[f"F{row}"].value for row in range(2, 5)]
        excel_requirements = [sheet[f"G{row}"].value for row in range(2, 5)]
        assert excel_priorities == ["Critical", "High", "Low"]
        assert excel_requirements[0] == "Prevent duplicate checkout charges"

        word = client.get("/export/scenarios/word", headers=headers)
        assert word.status_code == 200
        assert word.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        document = Document(BytesIO(word.content))
        exported_titles = [
            paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 1"
        ]
        assert exported_titles[:3] == [
            "1. Checkout blocks duplicate payment",
            "2. Login redirect regression",
            "3. Minor UI typo",
        ]
        full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert "Requirements: Prevent duplicate checkout charges" in full_text


def test_discussion_storage_and_enrichment_fallback() -> None:
    with TestClient(app) as client:
        headers = login(client)
        scenario = client.post(
            "/scenarios",
            headers=headers,
            json={
                "title": "Login redirects to homepage",
                "description": "Verify successful login lands on the homepage.",
                "steps": "Open login\nEnter valid credentials\nSubmit form",
                "expected_result": "User lands on the homepage",
                "priority": "high",
                "requirement_ids": [],
            },
        )
        assert scenario.status_code == 201

        topic = client.post("/topics", headers=headers, json={"title": "Password reset failures"})
        assert topic.status_code == 201
        topic_id = topic.json()["id"]

        created_message = client.post(
            f"/topics/{topic_id}/messages",
            headers=headers,
            json={
                "content": (
                    "Users cannot login after password reset and some requests "
                    "time out during authentication."
                )
            },
        )
        assert created_message.status_code == 201

        topics = client.get("/topics", headers=headers)
        assert topics.status_code == 200
        assert topics.json()[0]["enriched_message_count"] >= 1

        messages = client.get(f"/topics/{topic_id}/messages", headers=headers)
        assert messages.status_code == 200
        payload = messages.json()
        assert len(payload) == 1
        assert payload[0]["enriched_content"]
        assert payload[0]["enriched_content"].startswith("## Summary")
        assert "Fallback enrichment used" not in payload[0]["enriched_content"]
        assert "## Test Type Classification" not in payload[0]["enriched_content"]
        assert "## Test Ideas" in payload[0]["enriched_content"]
        assert " - functional" in payload[0]["enriched_content"]
        test_ideas_match = re.search(
            r"(?ms)^## Test Ideas\s*(.*?)(?=^##\s|\Z)",
            payload[0]["enriched_content"],
        )
        assert test_ideas_match is not None
        assert "[" not in test_ideas_match.group(1)
        assert "## Related Scenarios" in payload[0]["enriched_content"]
        assert "## QA Heuristics" in payload[0]["enriched_content"]
        assert payload[0]["enriched_content"].count("## QA Heuristics") == 1
        assert "Login redirects to homepage" in payload[0]["enriched_content"]


def test_scenario_suggestions_endpoint() -> None:
    with TestClient(app) as client:
        headers = login(client)
        topic = client.post("/topics", headers=headers, json={"title": "Login failures"})
        assert topic.status_code == 201
        topic_id = topic.json()["id"]

        message = client.post(
            f"/topics/{topic_id}/messages",
            headers=headers,
            json={"content": "Users log in successfully but do not reach the homepage."},
        )
        assert message.status_code == 201

        suggestions = client.post(f"/topics/{topic_id}/scenario-suggestions", headers=headers)
        assert suggestions.status_code == 200
        payload = suggestions.json()
        content = payload["content"]
        assert "## Scenario Suggestions" in content
        assert "Regression scenario" in content
        assert len(payload["scenarios"]) >= 1
        assert len(payload["scenarios"]) <= 3
        assert all(item["title"] for item in payload["scenarios"])
        assert all(item["description"] for item in payload["scenarios"])


def test_generate_and_save_scenarios_endpoint() -> None:
    with TestClient(app) as client:
        headers = login(client)
        topic = client.post("/topics", headers=headers, json={"title": "Login failures"})
        assert topic.status_code == 201
        topic_id = topic.json()["id"]

        message = client.post(
            f"/topics/{topic_id}/messages",
            headers=headers,
            json={"content": "Users log in successfully but do not reach the homepage."},
        )
        assert message.status_code == 201

        suggestions = client.post(f"/topics/{topic_id}/scenario-suggestions", headers=headers)
        assert suggestions.status_code == 200

        saved = client.post(
            f"/topics/{topic_id}/scenario-suggestions/save",
            headers=headers,
            json={"content": suggestions.json()["content"]},
        )
        assert saved.status_code == 201
        payload = saved.json()
        assert len(payload) >= 1
        assert all(item["title"] for item in payload)

        scenarios = client.get("/scenarios", headers=headers)
        assert scenarios.status_code == 200
        assert len(scenarios.json()) >= len(payload)


def test_requirement_hierarchy_linking_and_generation() -> None:
    with TestClient(app) as client:
        headers = login(client)
        parent = client.post(
            "/requirements",
            headers=headers,
            json={
                "title": "Checkout payment protection",
                "description": "The platform must protect the payment workflow from duplicate submissions.",
                "parent_id": None,
            },
        )
        assert parent.status_code == 201
        parent_id = parent.json()["id"]

        child = client.post(
            "/requirements",
            headers=headers,
            json={
                "title": "Reject duplicate submit clicks",
                "description": "A second checkout submit should be ignored while the first request is still processing.",
                "parent_id": parent_id,
            },
        )
        assert child.status_code == 201
        child_id = child.json()["id"]

        requirement_tree = client.get("/requirements", headers=headers)
        assert requirement_tree.status_code == 200
        payload = requirement_tree.json()
        assert any(item["id"] == parent_id for item in payload)
        assert payload[-1]["children"][0]["id"] == child_id or payload[0]["children"][0]["id"] == child_id

        linked_scenario = client.post(
            "/scenarios",
            headers=headers,
            json={
                "title": "Ignore second checkout submission",
                "description": "Covers the active processing duplicate-click case.",
                "steps": "Open checkout\nSubmit payment\nClick submit again immediately",
                "expected_result": "The duplicate click is ignored and only one payment is processed",
                "priority": "critical",
                "requirement_ids": [child_id],
            },
        )
        assert linked_scenario.status_code == 201
        assert linked_scenario.json()["requirements"][0]["id"] == child_id

        detail = client.get(f"/requirements/{child_id}", headers=headers)
        assert detail.status_code == 200
        detail_payload = detail.json()
        assert detail_payload["parent"]["id"] == parent_id
        assert any(item["title"] == "Ignore second checkout submission" for item in detail_payload["scenarios"])

        generated = client.post(f"/requirements/{child_id}/scenario-suggestions", headers=headers)
        assert generated.status_code == 200
        suggestion_payload = generated.json()
        assert "## Scenario Suggestions" in suggestion_payload["content"]
        assert len(suggestion_payload["scenarios"]) >= 1

        saved = client.post(
            f"/requirements/{child_id}/scenario-suggestions/save",
            headers=headers,
            json={"content": suggestion_payload["content"]},
        )
        assert saved.status_code == 201
        saved_payload = saved.json()
        assert len(saved_payload) >= 1
        assert all(
            any(requirement["id"] == child_id for requirement in item["requirements"])
            for item in saved_payload
        )


def test_role_permissions_enforced() -> None:
    with TestClient(app) as client:
        admin_headers = login(client)
        create_user(client, admin_headers, username="qa_user", role="qa")
        create_user(client, admin_headers, username="viewer_user", role="viewer")

        viewer_headers = login(client, username="viewer_user", password="Password123!")
        viewer_list = client.get("/scenarios", headers=viewer_headers)
        assert viewer_list.status_code == 200
        viewer_create = client.post(
            "/scenarios",
            headers=viewer_headers,
            json={
                "title": "Viewer should be blocked",
                "description": "",
                "steps": "",
                "expected_result": "",
                "priority": "low",
                "requirement_ids": [],
            },
        )
        assert viewer_create.status_code == 403

        qa_headers = login(client, username="qa_user", password="Password123!")
        qa_create = client.post(
            "/topics",
            headers=qa_headers,
            json={"title": "QA user can create discussion topics"},
        )
        assert qa_create.status_code == 201
        qa_user_admin_action = client.get("/users", headers=qa_headers)
        assert qa_user_admin_action.status_code == 403


def test_user_password_minimum_length_is_six() -> None:
    with TestClient(app) as client:
        admin_headers = login(client)

        too_short = client.post(
            "/users",
            headers=admin_headers,
            json={
                "username": "short_pw_user",
                "password": "12345",
                "role": "viewer",
            },
        )
        assert too_short.status_code == 422

        accepted = client.post(
            "/users",
            headers=admin_headers,
            json={
                "username": "six_pw_user",
                "password": "123456",
                "role": "viewer",
            },
        )
        assert accepted.status_code == 201


def test_parse_scenario_suggestions_rejects_tool_call_only_output() -> None:
    content = """```json
{ "name": "find_related_scenarios", "parameters": { "keyword": "login transfer homepage" } }
```"""

    assert parse_scenario_suggestions(content) == []


def test_parse_scenario_suggestions_rejects_field_fragments() -> None:
    content = """
## Scenario Suggestions

### Scenario 1: Login Transfer Issue
medium
Test Steps

-
Expected Result

-
### Priority: High
medium
Test Steps

-
Expected Result

-
### Steps:
medium
Test Steps

-
Expected Result

-
"""

    assert parse_scenario_suggestions(content) == []


def test_parse_scenario_suggestions_caps_results_to_three() -> None:
    content = """
## Scenario Suggestions

### One
- Description: Mainline scenario
- Priority: high
- Steps:
  - A
- Expected Result:
  - B

### Two
- Description: Backup scenario
- Priority: high
- Steps:
  - A
- Expected Result:
  - B

### Three
- Description: Third scenario
- Priority: high
- Steps:
  - A
- Expected Result:
  - B

### Four
- Description: Fourth scenario
- Priority: high
- Steps:
  - A
- Expected Result:
  - B
"""

    parsed = parse_scenario_suggestions(content)
    assert len(parsed) == 3
