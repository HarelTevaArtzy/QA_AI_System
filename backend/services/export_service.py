from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

from backend.models import Scenario


def _line_items(value: str) -> list[str]:
    return [line.strip() for line in value.splitlines() if line.strip()]


def _format_multiline(value: str) -> str:
    lines = _line_items(value)
    if not lines:
        return "-"
    return "\n".join(f"{index}. {line}" for index, line in enumerate(lines, start=1))


def build_scenarios_excel(scenarios: list[Scenario]) -> BytesIO:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "QA Scenarios"

    headers = [
        "ID",
        "Title",
        "Description",
        "Test Steps",
        "Expected Results",
        "Priority",
        "Created At",
    ]
    sheet.append(headers)

    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(fill_type="solid", fgColor="A63D40")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    for scenario in scenarios:
        sheet.append(
            [
                scenario.id,
                scenario.title,
                scenario.description,
                _format_multiline(scenario.steps),
                _format_multiline(scenario.expected_result),
                scenario.priority.title(),
                scenario.created_at.strftime("%Y-%m-%d %H:%M:%S"),
            ]
        )

    widths = {"A": 8, "B": 28, "C": 42, "D": 52, "E": 42, "F": 14, "G": 22}
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer


def build_scenarios_word(scenarios: list[Scenario]) -> BytesIO:
    document = Document()
    title = document.add_heading("QA Scenario Export", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph(
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for scenario in scenarios:
        document.add_heading(f"{scenario.id}. {scenario.title}", level=1)
        document.add_paragraph(f"Priority: {scenario.priority.title()}")
        if scenario.description:
            document.add_paragraph(scenario.description)

        document.add_paragraph("Test Steps", style="List Bullet")
        step_lines = _line_items(scenario.steps)
        if step_lines:
            for step in step_lines:
                document.add_paragraph(step, style="List Number")
        else:
            document.add_paragraph("No test steps documented yet.")

        document.add_paragraph("Expected Results", style="List Bullet")
        result_lines = _line_items(scenario.expected_result)
        if result_lines:
            for result in result_lines:
                document.add_paragraph(result, style="List Bullet 2")
        else:
            document.add_paragraph("No expected results documented yet.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
